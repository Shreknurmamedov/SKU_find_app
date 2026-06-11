from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from sku_audit.docx_catalog import import_docx_catalog


class DocxCatalogImportTests(unittest.TestCase):
    def test_imports_product_tables_with_category(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "catalog.docx"
            document = Document()
            document.add_paragraph("Каталог продукции Utake")
            document.add_paragraph("Дрель-шуруповерт (1)")
            table = document.add_table(rows=1, cols=4)
            for index, header in enumerate(["Бренд", "Модель", "Артикул", "Характеристики"]):
                table.rows[0].cells[index].text = header
            row = table.add_row().cells
            row[0].text = "Вихрь"
            row[1].text = "ДА-18Л-2К"
            row[2].text = "72/14/9"
            row[3].text = "Бренд: Вихрь; Напряжение: 18 В"
            document.save(path)

            products = import_docx_catalog(path)

            self.assertEqual(len(products), 1)
            self.assertEqual(products[0].brand_name, "Вихрь")
            self.assertEqual(products[0].category, "Дрель-шуруповерт")
            self.assertEqual(products[0].article_codes, ["72/14/9"])
            self.assertTrue(products[0].is_own_brand)


if __name__ == "__main__":
    unittest.main()
